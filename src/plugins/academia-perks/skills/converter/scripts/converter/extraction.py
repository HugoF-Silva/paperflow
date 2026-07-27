"""Extract DOCX papers as deterministic Markdown and JPEG figures."""
from __future__ import annotations

import io
import pathlib
import re
import shutil
import tempfile
from urllib.parse import unquote

import pypandoc
from docx import Document
from PIL import Image


_IMAGE_REFERENCE = re.compile(
    r'!\[(?P<alt>(?:\\.|[^\]\\])*)\]\((?P<target><[^>]+>|(?:\\.|[^)])+?)'
    r'(?P<title>\s+(?:"[^"]*"|\'[^\']*\'|\([^)]*\)))?\)'
    r'(?P<attrs>\{[^}\r\n]*\})?'
)
_RASTER_SUFFIXES = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}


def _resolve_media(target: str, raw_dir: pathlib.Path) -> pathlib.Path:
    value = unquote(target[1:-1] if target.startswith("<") else target)
    value = re.sub(r"\\([_*.{}\[\]()#! ])", r"\1", value)
    direct = pathlib.Path(value)
    raw_root = raw_dir.resolve()
    if direct.is_absolute() and not (
        direct.is_file() and direct.resolve().is_relative_to(raw_root)
    ):
        raise ValueError(f"Could not resolve extracted media reference: {target}")
    if direct.is_file() and direct.resolve().is_relative_to(raw_root):
        return direct.resolve()

    normalized = value.replace("\\", "/")
    matches = []
    for path in raw_dir.rglob("*"):
        resolved = path.resolve()
        if (
            path.is_file()
            and resolved.is_relative_to(raw_root)
            and normalized.endswith(path.relative_to(raw_dir).as_posix())
        ):
            matches.append(resolved)
    if len(matches) != 1:
        raise ValueError(f"Could not resolve extracted media reference: {target}")
    return matches[0]


def _save_jpeg(source: pathlib.Path, destination: pathlib.Path) -> None:
    if source.suffix.lower() == ".svg":
        import cairosvg

        image_source = io.BytesIO(cairosvg.svg2png(url=str(source)))
    elif source.suffix.lower() in _RASTER_SUFFIXES:
        image_source = source
    else:
        raise ValueError(f"Unsupported extracted media type: {source.suffix or source.name}")

    with Image.open(image_source) as image:
        rgba = image.convert("RGBA")
        rgb = Image.alpha_composite(
            Image.new("RGBA", rgba.size, "white"), rgba
        ).convert("RGB")
        rgb.save(
            destination,
            "JPEG",
            quality=95,
            subsampling=0,
            optimize=False,
            progressive=False,
        )


def _normalize_figures(
    markdown: str, raw_dir: pathlib.Path, figure_dir: pathlib.Path
) -> str:
    staged_dir = raw_dir / "normalized"
    staged_dir.mkdir()
    figures: dict[pathlib.Path, pathlib.Path] = {}

    def replace(reference: re.Match[str]) -> str:
        source = _resolve_media(reference.group("target"), raw_dir).resolve()
        if source not in figures:
            destination = staged_dir / f"figure-{len(figures) + 1:03d}.jpg"
            _save_jpeg(source, destination)
            figures[source] = destination
        final_path = figure_dir / figures[source].name
        title = reference.group("title") or ""
        attrs = reference.group("attrs") or ""
        return f'![{reference.group("alt")}](<{final_path.as_posix()}>{title}){attrs}'

    normalized = _IMAGE_REFERENCE.sub(replace, markdown)
    for staged in figures.values():
        staged.replace(figure_dir / staged.name)
    return normalized


def _with_headers_and_footers(docx_path: pathlib.Path, markdown: str) -> str:
    document = Document(docx_path)
    headers = [
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
        if paragraph.text
    ]
    footers = [
        paragraph.text
        for section in document.sections
        for paragraph in section.footer.paragraphs
        if paragraph.text
    ]
    parts = []
    if headers:
        parts.append(
            "--- DOCX HEADER ---\n"
            + "\n\n".join(headers)
            + "\n--- END DOCX HEADER ---"
        )
    parts.append(markdown.strip())
    if footers:
        parts.append(
            "--- DOCX FOOTER ---\n"
            + "\n\n".join(footers)
            + "\n--- END DOCX FOOTER ---"
        )
    return "\n\n".join(parts) + "\n"


def extract_paper(docx_path: pathlib.Path, workspace: pathlib.Path) -> str:
    """Return Markdown for ``docx_path`` and write JPEG figures in ``workspace``."""
    docx_path = pathlib.Path(docx_path).resolve()
    figure_dir = pathlib.Path(workspace) / "extracted_figures"
    if figure_dir.exists():
        shutil.rmtree(figure_dir)
    figure_dir.mkdir(parents=True)
    with tempfile.TemporaryDirectory(dir=figure_dir) as raw_dir:
        markdown = pypandoc.convert_file(
            str(docx_path),
            "markdown+tex_math_dollars",
            format="docx",
            extra_args=["--wrap=none", "--eol=lf", f"--extract-media={raw_dir}"],
        )
        markdown = _normalize_figures(markdown, pathlib.Path(raw_dir), figure_dir)
    return _with_headers_and_footers(docx_path, markdown)
