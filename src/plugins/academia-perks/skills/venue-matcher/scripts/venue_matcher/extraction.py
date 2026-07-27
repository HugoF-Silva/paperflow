"""Extract the full textual content of a .docx, headers/tables included, so the
author's stated target audience scope rides along in the text."""
from __future__ import annotations

import pathlib
from docx import Document


def extract_text(docx_path: pathlib.Path) -> str:
    doc = Document(str(docx_path))
    chunks: list[str] = []

    # Section headers/footers (where authors often write the target audience scope).
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                if para.text.strip():
                    chunks.append(para.text)

    # Body paragraphs.
    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)

    # Tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n\n".join(chunks)
